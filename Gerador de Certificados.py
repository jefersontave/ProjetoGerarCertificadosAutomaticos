import pandas as pd
from tkinter import *
from tkinter import messagebox, ttk
from docx import Document
from docx.shared import Pt

janela = Tk()
janela.title("Gerador de Certificados")

estilo = ttk.Style()
estilo.theme_use("alt")
estilo.configure(".", font="Arial 15", rowheight=30)

treeViewsDados = ttk.Treeview(janela, columns=(1,2,3,4,5,6), show="headings")

treeViewsDados.column("1", anchor=CENTER)
treeViewsDados.heading("1", text="CPF")

treeViewsDados.column("2", anchor=CENTER)
treeViewsDados.heading("2", text="NOME")

treeViewsDados.column("3", anchor=CENTER)
treeViewsDados.heading("3", text="RG")

treeViewsDados.column("4", anchor=CENTER)
treeViewsDados.heading("4", text="Data de Início")

treeViewsDados.column("5", anchor=CENTER)
treeViewsDados.heading("5", text="Data Final")

treeViewsDados.column("6", anchor=CENTER)
treeViewsDados.heading("6", text="Email")

treeViewsDados.grid(row=4, column=0 , columnspan=12, sticky="NSEW", pady=15)

def selecaoDuploClick(event):
    item =treeViewsDados.selection()
    for i in item:
        exibirCPF.delete(0, END)
        exibirNome.delete(0, END)
        exibirRG.delete(0, END)            #Limpando os campos
        exibirdataInicio.delete(0, END)
        exibirdataFim.delete(0, END)
        exibirEmail.delete(0, END)

        exibirCPF.insert(0, treeViewsDados.item(i, "values")[0])
        exibirNome.insert(0, treeViewsDados.item(i, "values")[1])
        exibirRG.insert(0, treeViewsDados.item(i, "values")[2])
        exibirdataInicio.insert(0, treeViewsDados.item(i, "values")[3])
        exibirdataFim.insert(0, treeViewsDados.item(i, "values")[4])
        exibirEmail.insert(0, treeViewsDados.item(i, "values")[5])


treeViewsDados.bind("<Double -1>", selecaoDuploClick)

#Abrindo o arquivo
dadosUsuarios =pd.read_excel("Dados.xlsx")
dadosUsuarios["Data Inicio"] = dadosUsuarios["Data Inicio"].astype(str)
dadosUsuarios["Data Fim"] = dadosUsuarios["Data Fim"].astype(str)


for linha in range(len(dadosUsuarios)):

    dataInicioAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
    dataInicioMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
    dataInicioDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

    dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

    dataFimAno = dadosUsuarios.iloc[linha, 4].split("-")[0]
    dataFimMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
    dataFimDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

    dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

    #Populando a Treeview com os dados do Excel
    treeViewsDados.insert("","end",values=(str(dadosUsuarios.iloc[linha, 0]), #CPF
                                           str(dadosUsuarios.iloc[linha, 1]), #NOME
                                           str(dadosUsuarios.iloc[linha, 2]), #RG
                                           str(dataInicioTratada), #Data de Início
                                           str(dataFimTratada), #Data Final
                                           str(dadosUsuarios.iloc[linha, 5])))#Email

cpf = Label(text="CPF: ", font="Arial 12")
cpf.grid(row=0,column=0,sticky="E", pady=15)

exibirCPF = Entry(font="Arial 12")
exibirCPF.grid(row=0,column=1,sticky="W",pady=15)

nome = Label(text="Nome: ", font="Arial 12")
nome.grid(row=0,column=2,sticky="E", pady=15)

exibirNome = Entry(font="Arial 12")
exibirNome.grid(row=0,column=3,sticky="W",pady=15)

RG = Label(text="RG: ", font="Arial 12")
RG.grid(row=0,column=4,sticky="E", pady=15)

exibirRG = Entry(font="Arial 12")
exibirRG.grid(row=0,column=5,sticky="W",pady=15)

dataInicio = Label(text="Data Inicio: ", font="Arial 12")
dataInicio.grid(row=0,column=6,sticky="E", pady=15)

exibirdataInicio = Entry(font="Arial 12")
exibirdataInicio.grid(row=0,column=7,sticky="W",pady=15)

dataFim = Label(text="Data Fim: ", font="Arial 12")
dataFim.grid(row=0,column=8,sticky="E", pady=15)

exibirdataFim = Entry(font="Arial 12")
exibirdataFim.grid(row=0,column=9,sticky="W",pady=15)

Email = Label(text="Email: ", font="Arial 12")
Email.grid(row=0,column=10,sticky="E", pady=15)

exibirEmail = Entry(font="Arial 12")
exibirEmail.grid(row=0,column=11,sticky="W",pady=15)

#------------------

def filtrar():
    for linha in range(len(dadosUsuarios)):
        todasLinhas = treeViewsDados.get_children()

        #deletando todas as linhas da treeview
        treeViewsDados.delete(*todasLinhas)

        if exibirCPF.get() == "":
            #Mudando o texto do botão
            botaoPesquisar.config(text="FILTRAR")
            
            for linha in range(len(dadosUsuarios)):
                dataInicioAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
                dataInicioMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
                dataInicioDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

                dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

                dataFimAno = dadosUsuarios.iloc[linha, 4].split("-")[0]
                dataFimMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
                dataFimDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

                dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

                # Populando a Treeview com os dados do Excel
                treeViewsDados.insert("", "end", values=(str(dadosUsuarios.iloc[linha, 0]),  # CPF
                                                         str(dadosUsuarios.iloc[linha, 1]),  # NOME
                                                         str(dadosUsuarios.iloc[linha, 2]),  # RG
                                                         str(dataInicioTratada),  # Data de Início
                                                         str(dataFimTratada),  # Data Final
                                                         str(dadosUsuarios.iloc[linha, 5])))  # Email
        else:
            botaoPesquisar.config(text="REMOVER FILTROS")
            for linha in range(len(dadosUsuarios)):
                if exibirCPF.get() == str(dadosUsuarios.iloc[linha, 0]):
                    dataInicioAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
                    dataInicioMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
                    dataInicioDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

                    dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

                    dataFimAno = dadosUsuarios.iloc[linha, 4].split("-")[0]
                    dataFimMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
                    dataFimDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

                    dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

                    treeViewsDados.insert("", "end", values=(str(dadosUsuarios.iloc[linha, 0]),  # CPF
                                                             str(dadosUsuarios.iloc[linha, 1]),  # NOME
                                                             str(dadosUsuarios.iloc[linha, 2]),  # RG
                                                             str(dataInicioTratada),  # Data de Início
                                                             str(dataFimTratada),  # Data Final
                                                             str(dadosUsuarios.iloc[linha, 5])))  # Email



botaoPesquisar = Button(text="PESQUISAR",font="Arial 15",border=3 , command=filtrar)
botaoPesquisar.grid(row=5, column=0, columnspan=4, sticky="NSEW", padx=20)

def gerarCertificado():
    #Abrindo o documento do word
    arquivoWord = Document("Certificado.docx")

    #Configurando o estilo
    estilo = arquivoWord.styles["Normal"]

    #Pegando os valores
    nomeAluno = exibirNome.get()
    dataInicio = exibirdataInicio.get()
    dataFim = exibirdataFim.get()
    instrutor = "Jeferson Borges Taveira"
    cpf_aluno = exibirCPF.get()
    rg_aluno = exibirRG.get()

    fraseMontada = f"A {nomeAluno} de CPF: {cpf_aluno} de RG: {rg_aluno} "\
              f"concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de RPA da " \
                   f"{dataInicio} a {dataFim}"

    for paragrafo in arquivoWord.paragraphs:
        if "@nome" in paragrafo.text:
            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calabri (Corpo)"
            fonte.size = Pt(24)
        if "@DataFim" in paragrafo.text:
            paragrafo.text = fraseMontada
            fonte = estilo.font
            fonte.name = "Calabri (Corpo)"
            fonte.size = Pt(24)
        if "Clevison Santos - Instrutor" in paragrafo.text:
            paragrafo.text = instrutor
            fonte = estilo.font
            fonte.name = "Calabri (Corpo)"
            fonte.size = Pt(24)

    #Salvando o certificado com o nome da pessoa
    caminhoSalvarCertificado = "C:\\Users\\jefer\\PycharmProjects\\RPA com python\\Certificados\\" + nomeAluno + ".docx"
    arquivoWord.save(caminhoSalvarCertificado)

    exibirCPF.delete(0, END)
    exibirNome.delete(0, END)
    exibirRG.delete(0, END)  # Limpando os campos
    exibirdataInicio.delete(0, END)
    exibirdataFim.delete(0, END)
    exibirEmail.delete(0, END)

    messagebox.showinfo("Mensagem", "Certificado gerado com sucesso")


botaoCertificado = Button(text="CERTIFICADO",font="Arial 15",border=3 , command=gerarCertificado)
botaoCertificado.grid(row=5, column=5, columnspan=4, sticky="NSEW", padx=20)

def gerarCertificadoParaTodos():
    #Pegando os valores da treeview por coluna
    for linha in treeViewsDados.get_children():
        coluna = treeViewsDados.item(linha)["values"]

        CPF = coluna[0]
        nome = coluna[1]
        RG = coluna[2]
        dataInicio = coluna[3]
        dataFim = coluna[4]
        instrutor = "Jeferson Borges Taveira"

        # Abrindo o documento do word
        arquivoWord = Document("Certificado.docx")

        # Configurando o estilo
        estilo = arquivoWord.styles["Normal"]

        fraseMontada = f"{nome} de CPF: {CPF} de RG: {RG} " \
                       f"concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de RPA da " \
                       f"{dataInicio} a {dataFim}"

        for paragrafo in arquivoWord.paragraphs:
            if "@nome" in paragrafo.text:
                paragrafo.text = nome
                fonte = estilo.font
                fonte.name = "Calabri (Corpo)"
                fonte.size = Pt(24)
            if "@DataFim" in paragrafo.text:
                paragrafo.text = fraseMontada
                fonte = estilo.font
                fonte.name = "Calabri (Corpo)"
                fonte.size = Pt(24)
            if "Clevison Santos - Instrutor" in paragrafo.text:
                paragrafo.text = instrutor
                fonte = estilo.font
                fonte.name = "Calabri (Corpo)"
                fonte.size = Pt(24)

        # Salvando o certificado com o nome da pessoa
        caminhoSalvarCertificado = "C:\\Users\\jefer\\PycharmProjects\\RPA com python\\Certificados\\" + nome + ".docx"
        arquivoWord.save(caminhoSalvarCertificado)

    messagebox.showinfo("Mensagem", "Certificados gerados com sucesso")

botaoCertificado = Button(text="CERTIFICADO PARA TODOS",font="Arial 15",border=3 , command=gerarCertificadoParaTodos)
botaoCertificado.grid(row=5, column=9, columnspan=4, sticky="NSEW", padx=20)
janela.mainloop()

